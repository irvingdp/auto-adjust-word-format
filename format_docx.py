#!/usr/bin/env python3
"""
Auto-adjust Word document format:
1. Clear headers & footers; remove tables containing GREENBRIER protocol banner text
2. Delete "No. of Samples" and "Banner" columns from tables
3. Split "Result/ Rating" into two columns: "Result" and "Rating"
4. Change all fonts to Tahoma, 10pt
5. Change orientation from landscape to portrait, auto-fit tables to window
6. If a table's estimated height exceeds one third of the page body, assign
   exact row heights whose total stays within one page (minus a small reserve)
"""

import copy
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from lxml import etree

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"


# ---------------------------------------------------------------------------
# XML-level helpers for table column manipulation
# ---------------------------------------------------------------------------

def get_grid_span(cell_el):
    tc_pr = cell_el.find(f"{W}tcPr")
    if tc_pr is not None:
        gs = tc_pr.find(f"{W}gridSpan")
        if gs is not None:
            return int(gs.get(f"{W}val", "1"))
    return 1


def set_grid_span(cell_el, span):
    tc_pr = cell_el.find(f"{W}tcPr")
    if tc_pr is None:
        tc_pr = etree.SubElement(cell_el, f"{W}tcPr")
        cell_el.insert(0, tc_pr)
    gs = tc_pr.find(f"{W}gridSpan")
    if span <= 1:
        if gs is not None:
            tc_pr.remove(gs)
    else:
        if gs is None:
            gs = etree.SubElement(tc_pr, f"{W}gridSpan")
        gs.set(f"{W}val", str(span))


def cell_text(cell_el):
    return "".join(t.text for t in cell_el.iter(f"{W}t") if t.text).strip()


def find_column_index(tbl_el, match_fn):
    """Return the grid-column index of the first header cell whose text satisfies *match_fn*, or -1."""
    rows = tbl_el.findall(f"{W}tr")
    if not rows:
        return -1
    col = 0
    for cell in rows[0].findall(f"{W}tc"):
        if match_fn(cell_text(cell)):
            return col
        col += get_grid_span(cell)
    return -1


def delete_column(tbl_el, col_idx):
    """Remove the grid column at *col_idx* from the table."""
    tbl_grid = tbl_el.find(f"{W}tblGrid")
    if tbl_grid is not None:
        grid_cols = tbl_grid.findall(f"{W}gridCol")
        if col_idx < len(grid_cols):
            tbl_grid.remove(grid_cols[col_idx])

    for tr in tbl_el.findall(f"{W}tr"):
        pos = 0
        for cell in tr.findall(f"{W}tc"):
            span = get_grid_span(cell)
            if pos <= col_idx < pos + span:
                if span == 1:
                    tr.remove(cell)
                else:
                    set_grid_span(cell, span - 1)
                break
            pos += span


def _clone_cell_shell(cell_el):
    """Create an empty <w:tc> that inherits tcPr (minus gridSpan) from *cell_el*."""
    new = etree.Element(f"{W}tc")
    old_pr = cell_el.find(f"{W}tcPr")
    if old_pr is not None:
        pr = copy.deepcopy(old_pr)
        gs = pr.find(f"{W}gridSpan")
        if gs is not None:
            pr.remove(gs)
        new.append(pr)
    p = etree.SubElement(new, f"{W}p")
    return new


def _copy_run_format(src_cell):
    """Return a deep-copied <w:rPr> from the first run of *src_cell*, or None."""
    for r in src_cell.iter(f"{W}r"):
        rpr = r.find(f"{W}rPr")
        if rpr is not None:
            return copy.deepcopy(rpr)
    return None


def split_result_rating(tbl_el):
    """Find the 'Result / Rating' column and split it into two columns."""
    col_idx = find_column_index(
        tbl_el,
        lambda txt: "result" in txt.lower() and "rating" in txt.lower(),
    )
    if col_idx < 0:
        return False

    # Widen tblGrid — each new column gets text+padding width
    result_w = (len("result") + 2) * DXA_PER_CHAR
    rating_w = (len("rating") + 2) * DXA_PER_CHAR
    tbl_grid = tbl_el.find(f"{W}tblGrid")
    if tbl_grid is not None:
        gcols = tbl_grid.findall(f"{W}gridCol")
        if col_idx < len(gcols):
            old = gcols[col_idx]
            old.set(f"{W}w", str(result_w))
            ng = etree.Element(f"{W}gridCol")
            ng.set(f"{W}w", str(rating_w))
            old.addnext(ng)

    rows = tbl_el.findall(f"{W}tr")
    header_rpr = None

    for ri, tr in enumerate(rows):
        pos = 0
        for cell in tr.findall(f"{W}tc"):
            span = get_grid_span(cell)
            if pos <= col_idx < pos + span:
                if span == 1:
                    new_cell = _clone_cell_shell(cell)
                    cell.addnext(new_cell)
                    if ri == 0:
                        header_rpr = _copy_run_format(cell)
                        # Keep only the first paragraph ("Result") in original cell
                        paras = cell.findall(f"{W}p")
                        for p in paras[1:]:
                            cell.remove(p)
                        # Set "Rating" in new cell with matching format
                        p = new_cell.find(f"{W}p")
                        if p is None:
                            p = etree.SubElement(new_cell, f"{W}p")
                        src_ppr = paras[0].find(f"{W}pPr") if paras else None
                        if src_ppr is not None:
                            p.insert(0, copy.deepcopy(src_ppr))
                        r = etree.SubElement(p, f"{W}r")
                        if header_rpr is not None:
                            r.insert(0, copy.deepcopy(header_rpr))
                        t = etree.SubElement(r, f"{W}t")
                        t.text = "Rating"
                else:
                    set_grid_span(cell, span + 1)
                break
            pos += span

    return True


# ---------------------------------------------------------------------------
# Header cleanup: consolidate runs, remove junk chars, center-align
# ---------------------------------------------------------------------------

# Tahoma 10pt ≈ 100 DXA per character.
# Column width = (len(header) + 2) * DXA_PER_CHAR  (1 char padding each side)
DXA_PER_CHAR = 100
FIXED_WIDTH_HEADERS = {"result", "rating", "country"}

HEADER_NAMES = {
    "evaluation", "citation / method", "citation/method",
    "criteria", "country", "result", "rating",
}


def _is_data_table(tbl_el):
    """Return True if the table looks like a data table (has known header names)."""
    rows = tbl_el.findall(f"{W}tr")
    if not rows:
        return False
    texts = set()
    for cell in rows[0].findall(f"{W}tc"):
        texts.add(cell_text(cell).lower().strip())
    return bool(texts & HEADER_NAMES)


def _clean_text(raw: str) -> str:
    """Normalise whitespace: collapse runs of spaces/newlines into single space, strip."""
    import re
    return re.sub(r"\s+", " ", raw).strip()


def clean_header_cells(doc):
    """For every data-table header row: consolidate runs into clean text, center-align."""
    count = 0
    for table in doc.tables:
        tbl_el = table._tbl
        if not _is_data_table(tbl_el):
            continue
        rows = tbl_el.findall(f"{W}tr")
        if not rows:
            continue

        row0 = rows[0]
        for cell in row0.findall(f"{W}tc"):
            raw = cell_text(cell)
            if not raw:
                continue
            clean = _clean_text(raw)

            # Grab the rPr from the first run as a formatting template
            base_rpr = _copy_run_format(cell)
            # Strip character-spacing overrides from template
            if base_rpr is not None:
                sp = base_rpr.find(f"{W}spacing")
                if sp is not None:
                    base_rpr.remove(sp)

            # Remove ALL existing paragraphs
            for p in list(cell.findall(f"{W}p")):
                cell.remove(p)

            # Build a single clean paragraph: center-aligned, one run
            p = etree.SubElement(cell, f"{W}p")
            ppr = etree.SubElement(p, f"{W}pPr")
            jc = etree.SubElement(ppr, f"{W}jc")
            jc.set(f"{W}val", "center")

            r = etree.SubElement(p, f"{W}r")
            if base_rpr is not None:
                r.insert(0, copy.deepcopy(base_rpr))
            t = etree.SubElement(r, f"{W}t")
            t.text = clean
            t.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")

        count += 1
    return count


# ---------------------------------------------------------------------------
# Font / size
# ---------------------------------------------------------------------------

def change_all_fonts(doc, name="Tahoma", size_pt=10):
    half_pt = str(size_pt * 2)
    body = doc.element.body

    for run in body.iter(f"{W}r"):
        rpr = run.find(f"{W}rPr")
        if rpr is None:
            rpr = etree.SubElement(run, f"{W}rPr")
            run.insert(0, rpr)

        rf = rpr.find(f"{W}rFonts")
        if rf is None:
            rf = etree.SubElement(rpr, f"{W}rFonts")
        rf.set(f"{W}ascii", name)
        rf.set(f"{W}hAnsi", name)
        rf.set(f"{W}eastAsia", name)
        rf.set(f"{W}cs", name)

        for tag in (f"{W}sz", f"{W}szCs"):
            el = rpr.find(tag)
            if el is None:
                el = etree.SubElement(rpr, tag)
            el.set(f"{W}val", half_pt)

    # Also update document default run properties (styles.xml)
    styles_el = doc.styles.element
    doc_defaults = styles_el.find(f"{W}docDefaults")
    if doc_defaults is None:
        doc_defaults = etree.SubElement(styles_el, f"{W}docDefaults")
        styles_el.insert(0, doc_defaults)
    rpr_default = doc_defaults.find(f"{W}rPrDefault")
    if rpr_default is None:
        rpr_default = etree.SubElement(doc_defaults, f"{W}rPrDefault")
    rpr = rpr_default.find(f"{W}rPr")
    if rpr is None:
        rpr = etree.SubElement(rpr_default, f"{W}rPr")

    rf = rpr.find(f"{W}rFonts")
    if rf is None:
        rf = etree.SubElement(rpr, f"{W}rFonts")
    rf.set(f"{W}ascii", name)
    rf.set(f"{W}hAnsi", name)
    rf.set(f"{W}eastAsia", name)
    rf.set(f"{W}cs", name)

    for tag in (f"{W}sz", f"{W}szCs"):
        el = rpr.find(tag)
        if el is None:
            el = etree.SubElement(rpr, tag)
        el.set(f"{W}val", str(size_pt * 2))


# ---------------------------------------------------------------------------
# Orientation & table auto-fit
# ---------------------------------------------------------------------------

def set_portrait(doc):
    for section in doc.sections:
        if section.orientation == WD_ORIENT.LANDSCAPE:
            w, h = section.page_width, section.page_height
            section.orientation = WD_ORIENT.PORTRAIT
            section.page_width = min(w, h)
            section.page_height = max(w, h)


def _compute_uniform_col_widths(doc):
    """Compute a single set of column widths for all data tables.

    Fixed-width columns (FIXED_WIDTH_HEADERS) get their defined widths.
    Remaining columns share the leftover space proportionally, using the
    majority pattern from the source tables as a base.
    """
    W_ = W
    # Collect gridCol widths from every data table
    all_widths = []
    for table in doc.tables:
        tbl_el = table._tbl
        rows = tbl_el.findall(f"{W_}tr")
        if not rows:
            continue
        headers = [cell_text(c).lower() for c in rows[0].findall(f"{W_}tc")]
        if "evaluation" not in headers:
            continue
        tbl_grid = tbl_el.find(f"{W_}tblGrid")
        if tbl_grid is None:
            continue
        gcols = [int(gc.get(f"{W_}w", "0")) for gc in tbl_grid.findall(f"{W_}gridCol")]
        all_widths.append(tuple(gcols))

    if not all_widths:
        return None

    # Use most-common pattern as base
    from collections import Counter
    base = list(Counter(all_widths).most_common(1)[0][0])
    num_cols = len(base)

    # Determine which grid indices are fixed-width (by checking first table header)
    sample_tbl = None
    for table in doc.tables:
        tbl_el = table._tbl
        rows = tbl_el.findall(f"{W_}tr")
        if not rows:
            continue
        headers = [cell_text(c).lower() for c in rows[0].findall(f"{W_}tc")]
        if "evaluation" in headers:
            sample_tbl = tbl_el
            break

    fixed = {}  # grid_col_index → dxa
    if sample_tbl is not None:
        pos = 0
        for hcell in sample_tbl.findall(f"{W_}tr")[0].findall(f"{W_}tc"):
            span = get_grid_span(hcell)
            txt = cell_text(hcell).lower()
            if txt in FIXED_WIDTH_HEADERS:
                w = (len(txt) + 2) * DXA_PER_CHAR
                for gi in range(pos, pos + span):
                    fixed[gi] = w
            pos += span

    # Apply fixed widths into the base; redistribute remaining space
    fixed_total = sum(fixed.values())
    flex_indices = [i for i in range(num_cols) if i not in fixed]
    flex_original_total = sum(base[i] for i in flex_indices)

    # Desired total = original total (keeps proportions the same, Word scales to 100%)
    original_total = sum(base)
    flex_target = original_total - fixed_total

    uniform = list(base)
    for gi, w in fixed.items():
        if gi < num_cols:
            uniform[gi] = w
    if flex_original_total > 0 and flex_target > 0:
        for i in flex_indices:
            uniform[i] = round(base[i] / flex_original_total * flex_target)

    return uniform


def autofit_tables_to_window(doc):
    """Set every table to 100 % page width with uniform column widths for data tables."""
    uniform = _compute_uniform_col_widths(doc)

    for tbl in doc.element.body.iter(f"{W}tbl"):
        tbl_pr = tbl.find(f"{W}tblPr")
        if tbl_pr is None:
            tbl_pr = etree.SubElement(tbl, f"{W}tblPr")
            tbl.insert(0, tbl_pr)

        # Table width = 100 %
        tw = tbl_pr.find(f"{W}tblW")
        if tw is None:
            tw = etree.SubElement(tbl_pr, f"{W}tblW")
        tw.set(f"{W}w", "5000")
        tw.set(f"{W}type", "pct")

        # Remove fixed layout
        layout = tbl_pr.find(f"{W}tblLayout")
        if layout is not None:
            tbl_pr.remove(layout)

        # Check if this is a data table
        rows = tbl.findall(f"{W}tr")
        is_data = False
        if rows:
            headers = [cell_text(c).lower() for c in rows[0].findall(f"{W}tc")]
            is_data = "evaluation" in headers

        if is_data and uniform is not None:
            # Apply uniform gridCol widths
            tbl_grid = tbl.find(f"{W}tblGrid")
            if tbl_grid is not None:
                gcols = tbl_grid.findall(f"{W}gridCol")
                for gi, gc in enumerate(gcols):
                    if gi < len(uniform):
                        gc.set(f"{W}w", str(uniform[gi]))

            # Apply matching tcW on every cell
            for tr in rows:
                pos = 0
                for tc in tr.findall(f"{W}tc"):
                    span = get_grid_span(tc)
                    tc_pr = tc.find(f"{W}tcPr")
                    if tc_pr is not None:
                        tcw = tc_pr.find(f"{W}tcW")
                        if tcw is not None:
                            cell_w = sum(uniform[pos:pos + span]) if pos + span <= len(uniform) else 0
                            tcw.set(f"{W}type", "dxa")
                            tcw.set(f"{W}w", str(cell_w))
                    pos += span
        else:
            # Non-data tables: just set auto
            for tc in tbl.iter(f"{W}tc"):
                tc_pr = tc.find(f"{W}tcPr")
                if tc_pr is not None:
                    tcw = tc_pr.find(f"{W}tcW")
                    if tcw is not None:
                        tcw.set(f"{W}type", "auto")
                        tcw.set(f"{W}w", "0")


# Estimated line height (twips) for Tahoma ~10pt after change_all_fonts.
_EST_LINE_TWIPS = 240

# When stretching a table toward one page, stay below the body height so grid
# lines, default cell padding, and layout rounding do not push the table past
# a single printed page.
_TABLE_STRETCH_RESERVE_TWIPS = 1440
_MIN_STRETCH_ROW_TWIPS = 50
# No single row may exceed this fraction of the table height budget (avoids
# one mega-row that still overflows visually with margins/borders).
_MAX_ROW_FRAC_OF_TABLE_TARGET = 0.40


def _section_body_height_twips(section):
    """Vertical space between top and bottom margins (twips)."""
    ph = section.page_height.twips
    tm = section.top_margin.twips if section.top_margin is not None else 1440
    bm = section.bottom_margin.twips if section.bottom_margin is not None else 1440
    return max(ph - tm - bm, 1)


def _estimated_row_height_twips(tr_el):
    tr_pr = tr_el.find(f"{W}trPr")
    if tr_pr is not None:
        th = tr_pr.find(f"{W}trHeight")
        if th is not None:
            val = th.get(f"{W}val")
            if val is not None:
                return max(int(val), 40)
    max_lines = 1
    for tc in tr_el.findall(f"{W}tc"):
        n = len(tc.findall(f"{W}p"))
        max_lines = max(max_lines, n)
    return max_lines * _EST_LINE_TWIPS


def _estimate_table_height_twips(tbl_el):
    return sum(_estimated_row_height_twips(tr) for tr in tbl_el.findall(f"{W}tr"))


def _set_tr_height_rule(tr_el, twips_val: int, rule: str):
    """rule: ``atLeast`` or ``exact`` (OOXML w:hRule)."""
    tr_pr = tr_el.find(f"{W}trPr")
    if tr_pr is None:
        tr_pr = etree.SubElement(tr_el, f"{W}trPr")
        tr_el.insert(0, tr_pr)
    th = tr_pr.find(f"{W}trHeight")
    if th is None:
        th = etree.SubElement(tr_pr, f"{W}trHeight")
    th.set(f"{W}val", str(max(twips_val, 1)))
    th.set(f"{W}hRule", rule)


def _distribute_row_heights_to_target(weights, target_total: int):
    """Integer row heights proportional to *weights*, sum == *target_total*.

    Each row is capped so one heavy content row cannot consume almost the full
    page (which tends to overflow in Word once padding and borders apply).

    Returns None if *target_total* is too small to give each row the minimum height.
    """
    n = len(weights)
    if n == 0 or target_total <= 0:
        return []
    total_w = sum(weights)
    max_row = max(
        _MIN_STRETCH_ROW_TWIPS,
        int(target_total * _MAX_ROW_FRAC_OF_TABLE_TARGET),
    )
    if total_w <= 0:
        base = max(target_total // n, _MIN_STRETCH_ROW_TWIPS)
        if base * n > target_total:
            return None
        return [min(base, max_row)] * n

    raw = [w * target_total / total_w for w in weights]
    heights = [
        max(_MIN_STRETCH_ROW_TWIPS, min(max_row, int(x)))
        for x in raw
    ]

    while sum(heights) > target_total:
        idx = max(range(n), key=lambda i: heights[i])
        if heights[idx] <= _MIN_STRETCH_ROW_TWIPS:
            return None
        heights[idx] -= 1

    deficit = target_total - sum(heights)
    fracs = sorted(
        range(n),
        key=lambda i: raw[i] - int(raw[i]),
        reverse=True,
    )
    k = 0
    guard = n * (target_total + 5)
    while deficit > 0 and k < guard:
        i = fracs[k % n]
        if heights[i] < max_row:
            heights[i] += 1
            deficit -= 1
        k += 1
    # If deficit remains, every row is at max_row — leave table slightly short
    # rather than exceeding one page.

    return heights


def stretch_tall_tables_to_page_body(doc):
    """Scale row heights when the table is moderately tall but under one page.

    Estimated height must be between one third of the page body and the full
    body. Row heights are set with ``exact`` so their sum stays at or below the
    printable body minus a small reserve, avoiding overflow past one page.
    """
    if not doc.sections:
        return 0
    section = doc.sections[0]
    body_h = _section_body_height_twips(section)
    threshold = body_h / 3.0
    target_cap = max(body_h - _TABLE_STRETCH_RESERVE_TWIPS, 1)
    stretched = 0

    for tbl in doc.element.body.iter(f"{W}tbl"):
        rows = tbl.findall(f"{W}tr")
        if not rows:
            continue
        row_heights = [_estimated_row_height_twips(tr) for tr in rows]
        total = sum(row_heights)
        n = len(rows)
        if total <= threshold or total >= body_h:
            continue
        if n * _MIN_STRETCH_ROW_TWIPS > target_cap:
            continue
        heights = _distribute_row_heights_to_target(row_heights, target_cap)
        if heights is None:
            continue
        for tr, h in zip(rows, heights):
            _set_tr_height_rule(tr, h, "exact")
        stretched += 1

    return stretched


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def fix_cell_paragraph_indents(doc):
    """Remove right indents and tab stops from all table-cell paragraphs.

    Landscape-origin right indents squeeze text in portrait columns,
    causing severe early wrapping.  Right indent in table cells is
    almost never intentional — remove it entirely.  Tab stops are also
    removed since they can push content beyond the visible cell area.
    """
    fixed = 0
    for tbl in doc.element.body.iter(f"{W}tbl"):
        for p in tbl.iter(f"{W}p"):
            ppr = p.find(f"{W}pPr")
            if ppr is None:
                continue

            ind = ppr.find(f"{W}ind")
            if ind is not None:
                right = ind.get(f"{W}right")
                if right and int(right) > 0:
                    ind.set(f"{W}right", "0")
                    fixed += 1

            tabs = ppr.find(f"{W}tabs")
            if tabs is not None:
                ppr.remove(tabs)

    return fixed


# Direct children of w:hdr / w:ftr that carry visible story content.
_HDR_FTR_BLOCK_TAGS = (f"{W}p", f"{W}tbl", f"{W}sdt")


def _clear_hdr_ftr_element(root_el):
    """Strip all block-level children from a header/footer root (w:hdr / w:ftr)."""
    for child in list(root_el):
        if child.tag in _HDR_FTR_BLOCK_TAGS:
            root_el.remove(child)


def clear_headers_and_footers(doc):
    """Remove all header and footer content from every section."""
    for section in doc.sections:
        for part in (
            section.header,
            section.first_page_header,
            section.even_page_header,
            section.footer,
            section.first_page_footer,
            section.even_page_footer,
        ):
            _clear_hdr_ftr_element(part._element)
            part.is_linked_to_previous = True


# Tables whose cell text (any cell, including nested content) contains this phrase are removed.
GREENBRIER_PROTOCOL_MARKER = "GREENBRIER INTERNATIONAL INC TEST PROTOCOL"


def _table_text_blob(tbl_el):
    return "".join((t.text or "") for t in tbl_el.iter(f"{W}t"))


def _collapse_whitespace_for_match(s: str) -> str:
    """Lowercase and strip all whitespace so line breaks between words still match.

    Word often splits banner titles across paragraphs; concatenating <w:t> runs
    then yields e.g. '...TEST' + 'PROTOCOL' with no space, which would not match
    the literal phrase 'TEST PROTOCOL'.
    """
    return "".join(s.split()).lower()


def remove_tables_containing_phrase(doc, phrase: str) -> int:
    """Drop any table that contains *phrase* in its text (case-insensitive).

    Whitespace is ignored for the substring test so titles split across
    paragraphs (e.g. ``TEST`` / ``PROTOCOL`` on separate lines) still match.

    When tables are nested, the **outermost** matching table is removed so the
    whole banner row (logos + text + page line) is stripped, not only an inner
    cell table.
    """
    needle = _collapse_whitespace_for_match(phrase)
    body = doc.element.body
    all_tbl = list(body.iter(f"{W}tbl"))
    matching = [
        tbl
        for tbl in all_tbl
        if needle in _collapse_whitespace_for_match(_table_text_blob(tbl))
    ]
    if not matching:
        return 0

    def has_matching_ancestor(t):
        return any(m is not t and m in t.iterancestors() for m in matching)

    to_remove = [t for t in matching if not has_matching_ancestor(t)]
    depth = lambda el: sum(1 for _ in el.iterancestors())
    removed = 0
    for tbl in sorted(to_remove, key=depth):
        parent = tbl.getparent()
        if parent is not None:
            parent.remove(tbl)
            removed += 1
    return removed


def process(input_path: str, output_path: str):
    doc = Document(input_path)

    # 0. Clear headers & footers
    clear_headers_and_footers(doc)
    print("  Cleared headers & footers")

    # 0b. Remove protocol banner tables
    n_banner = remove_tables_containing_phrase(doc, GREENBRIER_PROTOCOL_MARKER)
    print(f"  Removed {n_banner} table(s) containing protocol banner text")

    # 1. Delete columns
    cols_to_delete = ["No. of Samples"]
    for col_name in cols_to_delete:
        count = 0
        for table in doc.tables:
            idx = find_column_index(
                table._tbl, lambda t, cn=col_name: cn.lower() in t.lower()
            )
            if idx >= 0:
                delete_column(table._tbl, idx)
                count += 1
        print(f"  Deleted '{col_name}' from {count} table(s)")

    # 2. Split Result/Rating
    count = 0
    for table in doc.tables:
        if split_result_rating(table._tbl):
            count += 1
    print(f"  Split 'Result / Rating' in {count} table(s)")

    # 3. Clean header cells (consolidate text, remove junk, center-align)
    hdr_count = clean_header_cells(doc)
    print(f"  Cleaned & centered headers in {hdr_count} table(s)")

    # 4. Fonts
    change_all_fonts(doc, "Tahoma", 10)
    print("  Changed all fonts to Tahoma 10pt")

    # 5. Orientation + auto-fit
    set_portrait(doc)
    autofit_tables_to_window(doc)
    n_stretch = stretch_tall_tables_to_page_body(doc)
    print(
        "  Set portrait orientation & auto-fit tables to window"
        + (f"; stretched {n_stretch} tall table(s) to page height" if n_stretch else "")
    )

    # 6. Fix cell paragraph indents (remove right indents & tab stops)
    indent_fixed = fix_cell_paragraph_indents(doc)
    print(f"  Fixed {indent_fixed} paragraph(s) with excessive indent")

    Path(output_path).parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"  Saved → {output_path}")


if __name__ == "__main__":
    src = sys.argv[1] if len(sys.argv) > 1 else "source/原檔.docx"
    dst = sys.argv[2] if len(sys.argv) > 2 else "target/原檔_adjusted.docx"
    process(src, dst)
